from docx import Document
from docx.shared import Cm
import os
import io
class Comanda_text:
    def __init__(self, cliente:str, tel:str, servico:str,valor_total:str, sinal:str, restante:str,
                 data_entrada:str, data_retirada:str,comanda_id:str):
        self.cliente = cliente,
        self.tel = tel,
        self.servico = servico
        self.valor_total = valor_total
        self.sinal = sinal
        self.restante = restante
        self.data_entrada = data_entrada
        self.data_retirada = data_retirada
        self.comanda_id = comanda_id

    def set_custom_page_size(self,document, width_cm, height_cm):
        # Access the section element
        section = document.sections[0]

        # Set the page width and height
        section.page_width = Cm(width_cm)
        section.page_height = Cm(height_cm)

        # Optional: Set margins
        section.top_margin = Cm(.5)
        section.bottom_margin = Cm(.5)
        section.left_margin = Cm(.5)
        section.right_margin = Cm(.5)

    def text(self):
        doc = Document()
        self.set_custom_page_size(doc, 5.8, 20.99)

        #comanda_id:
        comanda_id = doc.add_paragraph()
        comanda_id.alignment = 2
        comanda_id.add_run(f'nº: {self.comanda_id}').bold=True

        #image:
        image_path = 'src/logo.png'
        im = doc.add_paragraph()
        im.alignment= 1
        imr = im.add_run()
        imr.add_picture(image_path, width=Cm(4))
        
        #infos:
        cl= doc.add_paragraph()
        cl.add_run('Cliente: ').bold=True
        cnome = self.cliente[0].replace('_', ' ')
        cl.add_run(cnome)
        tl= doc.add_paragraph()
        tl.add_run('Tel: ').bold=True
        tl.add_run(self.tel[0])
        se= doc.add_paragraph()
        se.add_run('Serviço: ').bold=True
        se.add_run(self.servico)
        vt= doc.add_paragraph()
        vt.add_run('Valor Total: ').bold=True
        vt.add_run(self.valor_total)
        si= doc.add_paragraph()
        si.add_run('Sinal: ').bold=True
        si.add_run(self.sinal)
        vr= doc.add_paragraph()
        vr.add_run('Valor Restante: ').bold=True
        vr.add_run(self.restante)
        de= doc.add_paragraph()
        de.add_run('Data entrada: ').bold=True
        de.add_run(self.data_entrada)
        dr= doc.add_paragraph()
        dr.add_run('Data retirada: ').bold=True
        dr.add_run(self.data_retirada)
        doc.add_paragraph(' ')
        ser= doc.add_paragraph()
        ser.add_run('O serviço deve ser retirado no prazo de 30 dias.').italic=True
        doc.add_paragraph(' ')
        duv= doc.add_paragraph()
        duv.add_run('Dúvidas e informações').italic=True

        doc.add_paragraph('(13) 99210-1642')

        #os.makedirs('data/docs',exist_ok=True)
        #doc.save(f'data/docs/comanda_{self.cliente[0]}_{self.comanda_id}.docx')
        print(f"Comanda criada")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
